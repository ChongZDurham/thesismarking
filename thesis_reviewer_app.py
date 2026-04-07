import json
import os
import queue
import re
import subprocess
import sys
import threading
import time
import traceback
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import requests
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    import pythoncom
    import win32com.client
except ImportError:  # pragma: no cover
    pythoncom = None
    win32com = None


APP_TITLE = "Thesis Reviewer / 论文评审助手"
SETTINGS_PATH = Path.home() / ".thesis_reviewer_settings.json"
SUPPORTED_SUFFIXES = {".doc", ".docx", ".txt", ".md", ".rtf", ".pdf"}
IS_WINDOWS = os.name == "nt"
IS_MACOS = sys.platform == "darwin"
MAX_REVIEW_ANNOTATIONS = 120
MAX_PARAGRAPH_REVIEWS = 48
SENTENCE_SCAN_BATCH_SIZE = 14
SENTENCE_SCAN_MAX_CHARS = 2200
PARAGRAPH_SCAN_BATCH_SIZE = 8
PARAGRAPH_SCAN_MAX_CHARS = 3200
CORE_ISSUE_TYPES = {"concept", "theory", "method", "evidence", "analysis", "ethics"}
SUBSTANTIVE_ISSUE_TYPES = {"concept", "literature", "theory", "method", "evidence", "analysis", "ethics"}
SURFACE_ISSUE_TYPES = {"language", "format"}
BARE_MINIMUM_STRENGTH_PATTERNS = (
    "基本结构",
    "结构完整",
    "框架完整",
    "章节齐全",
    "格式规范",
    "符合论文基本要求",
    "论文格式较为完整",
    "具有较完整的学位论文基本结构",
)
ARGUMENT_MARKERS = (
    "根据",
    "数据",
    "结果显示",
    "结果表明",
    "结果发现",
    "访谈",
    "问卷",
    "调查",
    "案例",
    "例如",
    "比如",
    "具体而言",
    "统计",
    "相关系数",
    "回归",
    "显著",
    "可见",
    "说明",
)
AI_JARGON_TERMS = (
    "多维",
    "多元协同",
    "系统性",
    "系统化",
    "机制建构",
    "路径优化",
    "路径构建",
    "协同治理",
    "赋能",
    "增能",
    "精准化",
    "全方位",
    "高质量",
    "实践范式",
    "治理效能",
    "复合支持",
    "联动机制",
    "闭环",
    "整合性",
    "嵌入式",
)
EMPTY_VALUE_PHRASES = (
    "具有重要意义",
    "具有现实意义",
    "具有理论意义",
    "具有实践意义",
    "实践意义",
    "理论意义",
    "提供参考",
    "提供借鉴",
    "有助于",
    "能够为",
    "奠定基础",
    "具有一定启示",
    "具有一定价值",
    "重要参考",
    "借鉴意义",
)
RESEARCH_QUESTION_MARKERS = (
    "研究问题",
    "研究目的",
    "本文旨在",
    "本研究旨在",
    "本研究拟",
    "本研究关注",
    "拟解决",
    "核心问题",
    "研究目标",
    "研究假设",
)
INTRO_SECTION_HINTS = ("绪论", "研究背景", "文献综述", "研究意义", "问题提出", "引言")
GENERIC_ANNOTATION_PATTERNS = (
    "该句包含的信息层次过多",
    "主干判断被多重补充成分打断",
    "该句抽象名词和专业术语连续堆叠",
    "该句并列成分过多",
    "该句把结论、依据和解释压在同一句里",
    "该句带有较强的价值判断色彩",
    "该句虽能表达基本意思，但表述仍偏概括",
    "该句措辞偏笼统，学术表达不够精准",
)
ANNOTATION_KEYWORD_STOPWORDS = {
    "本研究",
    "本文",
    "该句",
    "研究",
    "论文",
    "作者",
    "问题",
    "结果",
    "分析",
    "内容",
    "方面",
    "影响",
    "意义",
    "价值",
    "建议",
    "对象",
    "范围",
    "依据",
    "表述",
    "表达",
    "结论",
    "方法",
    "数据",
}
GENERIC_OVERALL_COMMENT_PATTERNS = (
    "该论文选题具有一定现实意义",
    "也能从社会工作视角切入问题",
    "达到基础写作要求、尚需较大幅度修改",
    "评分宜控制在合格到中等之间",
    "研究设计、统计处理与结论表达之间尚未充分对齐",
    "变量界定、研究工具、统计分析、结论强度和参考文献规范",
)
GENERIC_MAJOR_ISSUE_PATTERNS = (
    "论文整体基础一般",
    "方法、结果与结论之间的对应关系",
    "仍应重点强化",
)

SYSTEM_PROMPT = """
You are a careful thesis reviewer for student papers in social science and social work.
Write in Chinese. Be objective, specific, and constructive.
Do not invent evidence that is not in the paper.
Give a score out of 100.
Write an overall comment in about 300 Chinese words.
Use a strict grading standard:
- below 70 means obvious weakness or substantial revision still needed
- 70 means only "qualified/passable"
- 75-79 means moderate quality, still with clear weaknesses
- 80+ should be used only when the paper is clearly strong
- 85+ is reserved for genuinely excellent work
- 90+ should be very rare
Do not count bare-minimum structural completeness as a merit.
Generate 40-80 sentence-level annotations for a Word review report whenever the paper is long enough.
Each annotation must target a concrete sentence, claim, or very short passage from the paper, not a whole chapter.
The overall score must be broadly consistent with the dimension scores and major issues.
If method, evidence, or analysis is clearly weak, the overall score should not remain high.
Be alert to AI-style empty academic rhetoric: stacked professional nouns, generic value statements, and paragraphs that sound formal but do not actually argue.
Return JSON only.

JSON schema:
{
  "title": "paper title",
  "score": 74,
  "overall_comment_zh": "about 300 Chinese words",
  "summary_strengths": ["...", "..."],
  "major_issues": ["...", "..."],
  "dimension_scores": [
    {"name": "选题与价值", "score": 78, "comment": "..."},
    {"name": "文献与理论", "score": 74, "comment": "..."},
    {"name": "方法与数据", "score": 72, "comment": "..."},
    {"name": "分析与论证", "score": 71, "comment": "..."},
    {"name": "社会工作实践性", "score": 76, "comment": "..."},
    {"name": "语言与格式", "score": 73, "comment": "..."}
  ],
  "annotations": [
    {
      "section": "section or location",
      "quote": "short exact excerpt",
      "severity": "high|medium|low",
      "issue_type": "structure|concept|literature|theory|method|evidence|analysis|ethics|language|format",
      "comment": "diagnosis in Chinese",
      "suggestion": "revision advice in Chinese",
      "suggested_revision": "optional rewritten example in Chinese"
    }
  ],
  "paragraph_reviews": [
    {
      "section": "paragraph location",
      "quote": "short paragraph excerpt",
      "severity": "high|medium|low",
      "weakness_type": "no_claim|no_evidence|logic_jump|jargon_stack|empty_rhetoric|generic_conclusion|ai_generic_language|research_question_blur",
      "comment": "diagnosis in Chinese",
      "suggestion": "revision advice in Chinese",
      "suggested_revision": "optional rewritten example in Chinese"
    }
  ]
}
""".strip()

PARAGRAPH_SCAN_PROMPT = """
You are a strict paragraph-level reviewer for Chinese undergraduate theses in social science and social work.
Write in Chinese.
Do not praise.
Do not rescore the paper.
Judge whether each paragraph has real argumentation ability.
A strong paragraph usually has a concrete claim, some evidence or explanation, and a clear link between them.
Flag paragraphs that mainly stack professional jargon, give generic value statements, restate conclusions without support, or sound academic without saying anything specific.
Also flag paragraphs that look like generic AI writing, or paragraphs in the introduction/background/literature review that do not clearly serve a specific research question.
Return only the paragraphs that are weak, empty, or unconvincing.
Return JSON only.

JSON schema:
{
  "paragraph_reviews": [
    {
      "section": "provided paragraph label",
      "quote": "exact paragraph excerpt",
      "severity": "high|medium|low",
      "weakness_type": "no_claim|no_evidence|logic_jump|jargon_stack|empty_rhetoric|generic_conclusion|ai_generic_language|research_question_blur",
      "comment": "diagnosis in Chinese",
      "suggestion": "revision advice in Chinese",
      "suggested_revision": "optional rewritten example in Chinese"
    }
  ]
}
""".strip()

SENTENCE_SCAN_PROMPT = """
You are a strict sentence-level reviewer for Chinese undergraduate theses in social science and social work.
Write in Chinese.
Do not praise.
Do not rescore the paper.
Review the supplied sentences one by one.
Return annotations only for sentences that have a clear local problem.
A clear local problem includes overclaim, unsupported inference, vague concept use, missing method detail, weak evidence wording, logic jump, non-academic tone, grammar, or formatting.
Keep comments concrete and local. Prefer direct revision suggestions.
Do not default to saying that a sentence is too long.
Only use a length-based comment when sentence length itself clearly blocks meaning or argument.
When possible, diagnose the real problem more specifically: stacked abstract nouns, vague subject, missing evidence, too many parallel items, or a conclusion packed together with its basis.
Prefer content-specific critique over style critique.
At least 70% of the returned annotations should concern research content: concept definition, theory fit, literature use, method, evidence, analysis, ethics, or social-work professional logic.
Do not repeat stock comments. If you cannot name the concrete concept, variable, method, evidence, result, or professional issue that is problematic, skip that sentence.
Return JSON only.

JSON schema:
{
  "annotations": [
    {
      "section": "provided section label",
      "quote": "exact sentence from the list",
      "severity": "high|medium|low",
      "issue_type": "structure|concept|literature|theory|method|evidence|analysis|ethics|language|format",
      "comment": "diagnosis in Chinese",
      "suggestion": "revision advice in Chinese",
      "suggested_revision": "optional rewritten example in Chinese"
    }
  ]
}
""".strip()


@dataclass
class Config:
    paper: Path
    output_dir: Path
    api_base: str
    api_key: str
    model: str
    fallback_demo: bool
    annotated_copy: bool


class ReviewError(RuntimeError):
    pass


def read_text_with_fallback(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise ReviewError(f"Cannot decode file: {path}")


def clamp_score(value) -> int:
    try:
        value = int(value)
    except Exception:
        value = 0
    return max(0, min(100, value))


def safe_name(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", name).strip() or "review"


def compact_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def shorten_text(text: str, limit: int = 180) -> str:
    text = compact_whitespace(text)
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def dedupe_strings(items, limit: int = 8):
    cleaned = []
    seen = set()
    for item in items or []:
        text = compact_whitespace(item)
        if not text or text in seen:
            continue
        seen.add(text)
        cleaned.append(text)
        if limit and len(cleaned) >= limit:
            break
    return cleaned


def is_bare_minimum_strength(text: str) -> bool:
    text = compact_whitespace(text)
    if not text:
        return False
    return any(pattern in text for pattern in BARE_MINIMUM_STRENGTH_PATTERNS)


def dedupe_annotations(items, limit: int = MAX_REVIEW_ANNOTATIONS):
    cleaned = []
    seen = set()
    for item in items or []:
        if not isinstance(item, dict):
            continue
        section = compact_whitespace(item.get("section", "未定位")) or "未定位"
        quote = compact_whitespace(item.get("quote", ""))
        comment = compact_whitespace(item.get("comment", ""))
        suggestion = compact_whitespace(item.get("suggestion", ""))
        suggested_revision = compact_whitespace(item.get("suggested_revision", ""))
        if not quote and not comment and not suggestion:
            continue
        key = (section, quote, comment, suggestion)
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(
            {
                "section": section,
                "quote": quote,
                "severity": compact_whitespace(item.get("severity", "medium")).lower() or "medium",
                "issue_type": compact_whitespace(item.get("issue_type", "analysis")).lower() or "analysis",
                "comment": comment,
                "suggestion": suggestion,
                "suggested_revision": suggested_revision,
            }
        )
        if limit and len(cleaned) >= limit:
            break
    return cleaned


def dedupe_paragraph_reviews(items, limit: int = MAX_PARAGRAPH_REVIEWS):
    cleaned = []
    seen = set()
    for item in items or []:
        if not isinstance(item, dict):
            continue
        section = compact_whitespace(item.get("section", "未定位")) or "未定位"
        quote = compact_whitespace(item.get("quote", ""))
        comment = compact_whitespace(item.get("comment", ""))
        suggestion = compact_whitespace(item.get("suggestion", ""))
        suggested_revision = compact_whitespace(item.get("suggested_revision", ""))
        weakness_type = compact_whitespace(item.get("weakness_type", "empty_rhetoric")).lower() or "empty_rhetoric"
        if not quote and not comment and not suggestion:
            continue
        key = (section, quote, comment, suggestion, weakness_type)
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(
            {
                "section": section,
                "quote": quote,
                "severity": compact_whitespace(item.get("severity", "medium")).lower() or "medium",
                "weakness_type": weakness_type,
                "comment": comment,
                "suggestion": suggestion,
                "suggested_revision": suggested_revision,
            }
        )
        if limit and len(cleaned) >= limit:
            break
    return cleaned


def extract_annotation_keyphrases(text: str, limit: int = 8):
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9_-]{2,}|[\u4e00-\u9fff]{2,8}", compact_whitespace(text))
    picked = []
    seen = set()
    for token in sorted(tokens, key=len, reverse=True):
        token = token.strip()
        if token in seen:
            continue
        if token in ANNOTATION_KEYWORD_STOPWORDS:
            continue
        if len(token) < 2:
            continue
        seen.add(token)
        picked.append(token)
        if len(picked) >= limit:
            break
    return picked


def matched_generic_annotation_pattern(comment: str):
    comment = compact_whitespace(comment)
    for pattern in GENERIC_ANNOTATION_PATTERNS:
        if pattern in comment:
            return pattern
    return ""


def extract_json(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ReviewError("Model response did not contain valid JSON.")
    return json.loads(text[start : end + 1])


def is_com_busy(exc: Exception) -> bool:
    text = str(exc)
    return "被呼叫方拒绝接收呼叫" in text or "call was rejected by callee" in text


def call_with_retries(func, *args, retries: int = 12, delay: float = 0.5, **kwargs):
    last_exc = None
    for _ in range(retries):
        try:
            return func(*args, **kwargs)
        except Exception as exc:
            last_exc = exc
            if not is_com_busy(exc):
                raise
            time.sleep(delay)
    if last_exc:
        raise last_exc


def create_word_application():
    word = call_with_retries(win32com.client.DispatchEx, "Word.Application", retries=10, delay=0.6)
    try:
        word.Visible = False
    except Exception:
        pass
    try:
        word.DisplayAlerts = 0
    except Exception:
        pass
    try:
        word.ScreenUpdating = False
    except Exception:
        pass
    return word


def can_use_word_com() -> bool:
    return IS_WINDOWS and pythoncom is not None and win32com is not None


def read_docx_text(path: Path) -> str:
    if DocxDocument is None:
        raise ReviewError("python-docx is required to read .docx files on this platform.")
    try:
        document = DocxDocument(str(path))
    except Exception as exc:
        raise ReviewError(f"Failed to read .docx file: {exc}") from exc

    chunks = []
    for para in document.paragraphs:
        text = compact_whitespace(para.text)
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [compact_whitespace(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                chunks.append(" | ".join(cells))
    if not chunks:
        raise ReviewError("The .docx file did not contain readable text.")
    return "\n".join(chunks)


def read_pdf_text(path: Path) -> str:
    if PdfReader is None:
        raise ReviewError("pypdf is required to read PDF files on this platform.")
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ReviewError(f"Failed to open PDF file: {exc}") from exc

    chunks = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = text.strip()
        if text:
            chunks.append(text)
    if not chunks:
        raise ReviewError("The PDF file did not contain extractable text.")
    return "\n\n".join(chunks)


def read_with_textutil(path: Path) -> str:
    if not IS_MACOS:
        raise ReviewError("textutil is only available on macOS.")
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path.resolve())],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
    except FileNotFoundError as exc:
        raise ReviewError("macOS textutil was not found.") from exc
    except subprocess.CalledProcessError as exc:
        detail = (exc.stderr or exc.stdout or str(exc)).strip()
        raise ReviewError(f"Failed to read document with macOS textutil: {detail}") from exc

    text = result.stdout.strip()
    if not text:
        raise ReviewError("textutil did not return readable text.")
    return text


def open_path(path: Path):
    if IS_WINDOWS:
        os.startfile(str(path))
        return
    command = ["open", str(path)] if IS_MACOS else ["xdg-open", str(path)]
    subprocess.Popen(command)


def normalize_result(data: dict, title: str) -> dict:
    raw_score = clamp_score(data.get("score", 0))
    result = {
        "title": str(data.get("title") or title).strip(),
        "raw_score": raw_score,
        "score": raw_score,
        "overall_comment_zh": str(data.get("overall_comment_zh", "")).strip() or "未生成总体评语。",
        "summary_strengths": [],
        "major_issues": [],
        "dimension_scores": [],
        "annotations": [],
        "paragraph_reviews": [],
    }
    for item in data.get("summary_strengths", []):
        item = str(item).strip()
        if item and not is_bare_minimum_strength(item):
            result["summary_strengths"].append(item)
    for item in data.get("major_issues", []):
        item = str(item).strip()
        if item:
            result["major_issues"].append(item)
    for item in data.get("dimension_scores", []):
        if not isinstance(item, dict):
            continue
        result["dimension_scores"].append(
            {
                "name": str(item.get("name", "维度")).strip(),
                "score": clamp_score(item.get("score", 0)),
                "comment": str(item.get("comment", "")).strip(),
            }
        )
    for item in data.get("annotations", []):
        if not isinstance(item, dict):
            continue
        result["annotations"].append(
            {
                "section": str(item.get("section", "未定位")).strip(),
                "quote": str(item.get("quote", "")).strip(),
                "severity": str(item.get("severity", "medium")).strip().lower(),
                "issue_type": str(item.get("issue_type", "analysis")).strip().lower(),
                "comment": str(item.get("comment", "")).strip(),
                "suggestion": str(item.get("suggestion", "")).strip(),
                "suggested_revision": str(item.get("suggested_revision", "")).strip(),
            }
        )
    for item in data.get("paragraph_reviews", []):
        if not isinstance(item, dict):
            continue
        result["paragraph_reviews"].append(
            {
                "section": str(item.get("section", "未定位")).strip(),
                "quote": str(item.get("quote", "")).strip(),
                "severity": str(item.get("severity", "medium")).strip().lower(),
                "weakness_type": str(item.get("weakness_type", "empty_rhetoric")).strip().lower(),
                "comment": str(item.get("comment", "")).strip(),
                "suggestion": str(item.get("suggestion", "")).strip(),
                "suggested_revision": str(item.get("suggested_revision", "")).strip(),
            }
        )
    result["summary_strengths"] = dedupe_strings(result["summary_strengths"], limit=6)
    result["major_issues"] = dedupe_strings(result["major_issues"], limit=8)
    result["annotations"] = dedupe_annotations(result["annotations"])
    result["paragraph_reviews"] = dedupe_paragraph_reviews(result["paragraph_reviews"])
    return result


class ReviewEngine:
    def __init__(self, logger):
        self.logger = logger

    def log(self, message: str):
        self.logger(message)

    def get_dimension_score(self, result: dict, keywords):
        keywords = tuple(keywords)
        for item in result.get("dimension_scores", []):
            name = compact_whitespace(item.get("name", ""))
            if any(keyword in name for keyword in keywords):
                return clamp_score(item.get("score", 0))
        return None

    def annotation_specificity_score(self, item: dict):
        issue_type = compact_whitespace(item.get("issue_type", "")).lower()
        quote = compact_whitespace(item.get("quote", ""))
        comment = compact_whitespace(item.get("comment", ""))
        suggestion = compact_whitespace(item.get("suggestion", ""))
        combined = f"{comment} {suggestion}"
        score = 0
        if issue_type in SUBSTANTIVE_ISSUE_TYPES:
            score += 4
        if issue_type in SURFACE_ISSUE_TYPES:
            score -= 1
        if matched_generic_annotation_pattern(comment):
            score -= 4
        overlap = sum(1 for token in extract_annotation_keyphrases(quote) if token and token in combined)
        score += min(overlap, 3)
        if any(marker in combined for marker in ("变量", "量表", "理论", "文献", "样本", "问卷", "访谈", "结果", "统计", "中介", "社会工作", "服务对象", "研究问题")):
            score += 1
        if len(comment) >= 18:
            score += 1
        return score

    def prioritize_annotations(self, annotations):
        items = []
        generic_pattern_counts = {}
        for item in dedupe_annotations(annotations):
            updated = dict(item)
            pattern = matched_generic_annotation_pattern(updated.get("comment", ""))
            if pattern:
                generic_pattern_counts[pattern] = generic_pattern_counts.get(pattern, 0) + 1
                if generic_pattern_counts[pattern] > 1:
                    continue
            spec = self.annotation_specificity_score(updated)
            issue_type = compact_whitespace(updated.get("issue_type", "")).lower()
            if issue_type in SURFACE_ISSUE_TYPES and spec < 2:
                continue
            if spec < 0:
                continue
            updated["_spec"] = spec
            items.append(updated)

        severity_order = {"high": 0, "medium": 1, "low": 2}
        substantive = [item for item in items if item.get("issue_type") in SUBSTANTIVE_ISSUE_TYPES]
        surface = [item for item in items if item.get("issue_type") not in SUBSTANTIVE_ISSUE_TYPES]

        substantive.sort(key=lambda item: (severity_order.get(item.get("severity"), 3), -item.get("_spec", 0)))
        surface.sort(key=lambda item: (severity_order.get(item.get("severity"), 3), -item.get("_spec", 0)))

        selected = list(substantive)
        if substantive:
            max_surface = max(2, min(5, len(substantive) // 3 + 1))
            selected.extend(surface[:max_surface])
        else:
            selected.extend(surface[:3])

        selected.sort(key=lambda item: (severity_order.get(item.get("severity"), 3), 0 if item.get("issue_type") in SUBSTANTIVE_ISSUE_TYPES else 1, -item.get("_spec", 0)))
        cleaned = []
        for item in selected[:MAX_REVIEW_ANNOTATIONS]:
            stripped = dict(item)
            stripped.pop("_spec", None)
            cleaned.append(stripped)
        return dedupe_annotations(cleaned)

    def needs_overall_comment_refresh(self, comment: str, title: str):
        comment = compact_whitespace(comment)
        if not comment:
            return True
        if any(pattern in comment for pattern in GENERIC_OVERALL_COMMENT_PATTERNS):
            return True
        title_tokens = extract_annotation_keyphrases(title, limit=5)
        if title_tokens and not any(token in comment for token in title_tokens):
            generic_hits = sum(1 for pattern in GENERIC_OVERALL_COMMENT_PATTERNS if pattern in comment)
            if generic_hits >= 2:
                return True
        return False

    def strip_terminal_punct(self, text: str):
        return compact_whitespace(text).rstrip("。；;，, ")

    def is_generic_major_issue(self, text: str):
        text = compact_whitespace(text)
        return any(pattern in text for pattern in GENERIC_MAJOR_ISSUE_PATTERNS)

    def issue_type_label(self, issue_type: str):
        mapping = {
            "concept": "概念界定",
            "literature": "文献使用",
            "theory": "理论框架",
            "method": "研究方法",
            "evidence": "证据支撑",
            "analysis": "分析论证",
            "ethics": "研究伦理",
            "language": "语言表达",
            "format": "格式规范",
        }
        return mapping.get(issue_type, issue_type or "内容问题")

    def compress_issue_phrase(self, text: str):
        text = self.strip_terminal_punct(text)
        shortcuts = (
            ("AI 化通用写法", "AI 化通用语言"),
            ("模板化学术语言", "模板化学术语言"),
            ("原创表达不足", "原创表达不足"),
            ("研究问题界定不清", "研究问题界定不清"),
            ("研究背景、文献梳理与研究目标之间的服务关系不足", "背景与综述未服务核心问题"),
            ("工具说明", "研究工具说明不足"),
            ("结果支撑不足", "结果支撑不足"),
            ("概念口径前后不够一致", "概念口径不一致"),
        )
        for pattern, short in shortcuts:
            if pattern in text:
                return short
        return text

    def weakness_type_label(self, weakness_type: str):
        mapping = {
            "ai_generic_language": "AI 化通用语言",
            "research_question_blur": "研究问题不清",
            "jargon_stack": "术语堆砌",
            "empty_rhetoric": "空泛价值判断",
            "no_claim": "缺少明确论点",
            "no_evidence": "缺少论据支撑",
            "logic_jump": "论证跳步",
            "generic_conclusion": "结论泛化",
        }
        return mapping.get(weakness_type, weakness_type or "段落问题")

    def compose_specific_overall_comment(self, result: dict, title: str, text: str):
        title_text = compact_whitespace(title) or "该论文"
        strengths = [self.strip_terminal_punct(item) for item in result.get("summary_strengths", []) if compact_whitespace(item)]
        major_issues = [self.strip_terminal_punct(item) for item in result.get("major_issues", []) if compact_whitespace(item)]
        dims = sorted(result.get("dimension_scores", []), key=lambda item: clamp_score(item.get("score", 0)))
        paragraph_reviews = result.get("paragraph_reviews", [])
        annotations = result.get("annotations", [])
        specific_issues = [self.compress_issue_phrase(item) for item in major_issues if not self.is_generic_major_issue(item)]

        weakest_dims = [item["name"] for item in dims[:2] if item.get("name")]
        dominant_issue_types = []
        seen_issue_types = set()
        for item in annotations:
            issue_type = self.issue_type_label(item.get("issue_type", ""))
            if issue_type in seen_issue_types:
                continue
            seen_issue_types.add(issue_type)
            dominant_issue_types.append(issue_type)
            if len(dominant_issue_types) >= 3:
                break

        paragraph_flags = []
        seen_flags = set()
        for item in paragraph_reviews:
            label = self.weakness_type_label(item.get("weakness_type", ""))
            if label in seen_flags:
                continue
            seen_flags.add(label)
            paragraph_flags.append(label)
            if len(paragraph_flags) >= 3:
                break

        sentences = [f"论文围绕《{title_text}》展开。"]
        if strengths:
            sentences.append("相对可取之处在于" + "；".join(strengths[:2]) + "。")
        if not result.get("research_question_marker_found", True):
            sentences.append("当前最核心的缺陷是研究问题界定不清，研究背景和文献梳理没有稳定服务于一个明确的研究目标。")
        elif specific_issues:
            sentences.append("目前更突出的核心问题包括" + "；".join(specific_issues[:2]) + "。")
        if weakest_dims:
            sentences.append("从评分结构看，" + "和".join(weakest_dims) + "是当前最需要优先补强的部分。")
        if dominant_issue_types:
            sentences.append("句子级批注主要集中在" + "、".join(dominant_issue_types) + "，说明问题并不只是形式层面，而是涉及论证内容本身。")
        if paragraph_flags:
            sentences.append("段落层面还识别出" + "、".join(paragraph_flags) + "等现象，提示部分文字存在空泛化或目标不聚焦的问题。")
        if not specific_issues and not paragraph_flags and dominant_issue_types:
            sentences.append("就内容层面看，当前短板主要不是格式，而是" + "、".join(dominant_issue_types) + "尚未真正支撑起全文的核心论证。")
        if result.get("score_calibration_note"):
            sentences.append("本次评分也已结合这些实质性问题做过校准，不会仅因结构完整或语言看起来顺畅而给高分。")
        if specific_issues:
            sentences.append("建议下一轮修改先围绕" + "、".join(specific_issues[:3]) + "逐项处理，再统一修整语言和格式。")
        elif dominant_issue_types:
            sentences.append("建议下一轮修改优先处理" + "、".join(dominant_issue_types[:3]) + "，再回头修整语言和格式层面的细节。")
        else:
            sentences.append("建议下一轮修改继续围绕研究问题、论证链条和证据支撑做实质性补强。")
        return "".join(sentences)

    def text_has_clear_research_question(self, text: str) -> bool:
        text = compact_whitespace(text)
        return any(marker in text for marker in RESEARCH_QUESTION_MARKERS)

    def augment_major_issues_from_signals(self, result: dict, text: str):
        paragraph_reviews = result.get("paragraph_reviews", [])
        major_issues = list(result.get("major_issues", []))
        ai_generic_count = sum(1 for item in paragraph_reviews if item.get("weakness_type") == "ai_generic_language")
        research_blur_count = sum(1 for item in paragraph_reviews if item.get("weakness_type") == "research_question_blur")
        has_research_question_marker = self.text_has_clear_research_question(text)

        if ai_generic_count >= 1 or sum(1 for item in paragraph_reviews if item.get("weakness_type") == "jargon_stack") >= 2:
            major_issues.append("部分段落呈现明显的 AI 化通用写法或模板化学术语言，原创表达不足，不宜给高分。")
        if research_blur_count >= 1 or not has_research_question_marker:
            major_issues.append("研究问题界定不清，研究背景、文献梳理与研究目标之间的服务关系不足，核心目标不够聚焦。")

        result["major_issues"] = dedupe_strings(major_issues, limit=8)
        result["research_question_marker_found"] = has_research_question_marker
        return result

    def calibrate_review_result(self, result: dict):
        score = clamp_score(result.get("score", 0))
        raw_score = clamp_score(result.get("raw_score", score))
        dimensions = result.get("dimension_scores", [])
        annotations = result.get("annotations", [])
        paragraph_reviews = result.get("paragraph_reviews", [])
        major_issues = result.get("major_issues", [])

        avg_dim = None
        if dimensions:
            avg_dim = round(sum(clamp_score(item.get("score", 0)) for item in dimensions) / len(dimensions))
            score = min(score, avg_dim)

        high_count = sum(1 for item in annotations if item.get("severity") == "high")
        medium_count = sum(1 for item in annotations if item.get("severity") == "medium")
        weak_paragraph_count = len(paragraph_reviews)
        high_paragraph_count = sum(1 for item in paragraph_reviews if item.get("severity") == "high")
        jargon_paragraph_count = sum(1 for item in paragraph_reviews if item.get("weakness_type") in {"jargon_stack", "empty_rhetoric"})
        ai_generic_count = sum(1 for item in paragraph_reviews if item.get("weakness_type") == "ai_generic_language")
        ai_generic_high_count = sum(
            1 for item in paragraph_reviews if item.get("weakness_type") == "ai_generic_language" and item.get("severity") == "high"
        )
        research_blur_count = sum(1 for item in paragraph_reviews if item.get("weakness_type") == "research_question_blur")
        core_high_count = sum(
            1
            for item in annotations
            if item.get("severity") == "high" and item.get("issue_type") in CORE_ISSUE_TYPES
        )
        core_issue_types = {
            item.get("issue_type")
            for item in annotations
            if item.get("issue_type") in CORE_ISSUE_TYPES and item.get("severity") in {"high", "medium"}
        }
        low_dims = [item for item in dimensions if clamp_score(item.get("score", 0)) <= 69]
        very_low_dims = [item for item in dimensions if clamp_score(item.get("score", 0)) <= 64]
        method_score = self.get_dimension_score(result, ("方法", "数据"))
        analysis_score = self.get_dimension_score(result, ("分析", "论证"))
        language_score = self.get_dimension_score(result, ("语言", "格式"))

        caps = []
        if high_count >= 1:
            caps.append(84)
        if high_count >= 2 or len(major_issues) >= 4:
            caps.append(79)
        if high_count >= 3 or len(low_dims) >= 3 or len(core_issue_types) >= 3:
            caps.append(74)
        if high_count >= 4 or core_high_count >= 2 or len(major_issues) >= 6 or len(very_low_dims) >= 2:
            caps.append(69)
        if high_count >= 5 or core_high_count >= 3 or (avg_dim is not None and avg_dim <= 64):
            caps.append(64)
        if weak_paragraph_count >= 3:
            caps.append(76)
        if jargon_paragraph_count >= 1:
            caps.append(78)
        if weak_paragraph_count >= 5 or jargon_paragraph_count >= 2:
            caps.append(72)
        if high_paragraph_count >= 2:
            caps.append(68)
        if ai_generic_count >= 1:
            caps.append(74)
        if ai_generic_count >= 2 or ai_generic_high_count >= 1:
            caps.append(70)
        if ai_generic_count >= 4 or ai_generic_high_count >= 2:
            caps.append(66)
        if research_blur_count >= 1:
            caps.append(76)
        if research_blur_count >= 2 or not result.get("research_question_marker_found", True):
            caps.append(72)
        if method_score is not None and analysis_score is not None and method_score <= 65 and analysis_score <= 65:
            caps.append(69)
        if method_score is not None and analysis_score is not None and method_score <= 60 and analysis_score <= 60:
            caps.append(64)

        score -= min(max(0, medium_count - 8), 3)
        if language_score is not None and language_score <= 65:
            score -= 1
        if caps:
            score = min(score, min(caps))

        score = clamp_score(score)
        result["score"] = score
        if score != raw_score:
            note_parts = [f"原始模型分数 {raw_score}", f"校准后分数 {score}"]
            if avg_dim is not None:
                note_parts.append(f"维度均分 {avg_dim}")
            if high_count:
                note_parts.append(f"高严重度问题 {high_count} 条")
            if weak_paragraph_count:
                note_parts.append(f"论证薄弱段落 {weak_paragraph_count} 段")
            if ai_generic_count:
                note_parts.append(f"AI 化语言段落 {ai_generic_count} 段")
            if research_blur_count or not result.get("research_question_marker_found", True):
                note_parts.append("研究问题聚焦不足")
            if len(major_issues):
                note_parts.append(f"主要问题 {len(major_issues)} 项")
            result["score_calibration_note"] = "；".join(note_parts)
        else:
            result["score_calibration_note"] = ""
        return result

    def review(self, cfg: Config) -> dict:
        self.log("Loading paper...")
        title, text = self.load_paper(cfg.paper)
        self.log("Paper loaded.")
        prompt_text = self.trim_for_model(text)
        review_mode = "api"

        try:
            if cfg.api_key.strip():
                self.log("Calling review model...")
                raw = self.call_model(title, prompt_text, cfg)
                try:
                    self.log("Running paragraph argumentation scan...")
                    paragraph_reviews = self.generate_paragraph_argumentation_reviews(title, text, cfg)
                    raw["paragraph_reviews"] = dedupe_paragraph_reviews(paragraph_reviews + list(raw.get("paragraph_reviews", [])))
                except Exception as exc:
                    self.log(f"Paragraph argumentation scan skipped: {exc}")
                try:
                    self.log("Running sentence-level scan for detailed annotations...")
                    detailed_annotations = self.generate_sentence_level_annotations(title, text, cfg)
                    raw["annotations"] = dedupe_annotations(detailed_annotations + list(raw.get("annotations", [])))
                except Exception as exc:
                    self.log(f"Sentence-level scan skipped: {exc}")
            elif cfg.fallback_demo:
                self.log("No API key found. Using demo heuristic review.")
                review_mode = "demo"
                raw = self.demo_review(title, text)
            else:
                raise ReviewError("API key is required unless demo review is enabled.")
        except Exception as exc:
            if cfg.fallback_demo:
                self.log(f"Model review failed. Falling back to demo review: {exc}")
                review_mode = "demo"
                raw = self.demo_review(title, text)
            else:
                raise

        result = normalize_result(raw, title)
        result["review_mode"] = review_mode
        result["annotations"] = self.prioritize_annotations(result.get("annotations", []))
        result = self.augment_major_issues_from_signals(result, text)
        result = self.calibrate_review_result(result)
        if review_mode == "demo" or self.needs_overall_comment_refresh(result.get("overall_comment_zh", ""), title):
            result["overall_comment_zh"] = self.compose_specific_overall_comment(result, title, text)
        if result.get("score_calibration_note"):
            self.log(result["score_calibration_note"])
        result["outputs"] = self.write_outputs(cfg.paper, cfg.output_dir, result, cfg.annotated_copy)
        return result

    def load_paper(self, path: Path):
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            raise ReviewError(f"Unsupported file type: {path.suffix}")
        if suffix in {".txt", ".md"}:
            text = read_text_with_fallback(path)
        elif suffix == ".docx":
            if DocxDocument is not None:
                text = read_docx_text(path)
            elif can_use_word_com():
                text = self.read_with_word(path)
            elif IS_MACOS:
                text = read_with_textutil(path)
            else:
                raise ReviewError("Reading .docx files requires python-docx on this platform.")
        elif suffix == ".pdf":
            text = read_pdf_text(path)
        elif suffix in {".doc", ".rtf"}:
            if can_use_word_com():
                text = self.read_with_word(path)
            elif IS_MACOS:
                text = read_with_textutil(path)
            else:
                raise ReviewError("Reading .doc/.rtf files requires Word on Windows or textutil on macOS.")
        else:
            text = self.read_with_word(path)
        text = text.replace("\r", "\n").replace("\x07", " ").replace("\x0c", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        title = self.extract_title(text, path)
        return title, text

    def extract_title(self, text: str, path: Path):
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        for line in lines[:30]:
            if len(line) < 8:
                continue
            if any(x in line for x in ("本科毕业论文", "诚信声明", "授权说明", "学生姓名", "指导教师", "提交日期")):
                continue
            return line
        return path.stem

    def trim_for_model(self, text: str, limit: int = 40000):
        if len(text) <= limit:
            return text
        return text[:32000] + "\n\n[...truncated due to length...]\n\n" + text[-8000:]

    def request_json_completion(self, cfg: Config, system_prompt: str, user_prompt: str, timeout: int = 240):
        base = cfg.api_base.strip().rstrip("/")
        if not base:
            raise ReviewError("API base URL is empty.")
        endpoint = base if base.endswith("/chat/completions") else base + "/chat/completions"
        payload = {
            "model": cfg.model.strip(),
            "temperature": 0.2,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "response_format": {"type": "json_object"},
        }
        headers = {
            "Authorization": f"Bearer {cfg.api_key.strip()}",
            "Content-Type": "application/json",
        }
        res = requests.post(endpoint, headers=headers, json=payload, timeout=timeout)
        res.raise_for_status()
        content = res.json()["choices"][0]["message"]["content"]
        if isinstance(content, list):
            content = "".join(part.get("text", "") for part in content if isinstance(part, dict))
        if not isinstance(content, str):
            raise ReviewError("Unexpected model response format.")
        return extract_json(content)

    def call_model(self, title: str, text: str, cfg: Config):
        user_prompt = f"""
Paper title: {title}

Task:
1. Review this thesis in Chinese.
2. Give a score out of 100.
3. Write an overall comment in about 300 Chinese words.
4. Generate 40-80 sentence-level annotations if the paper is long enough.
5. Focus on structure, concepts, literature, methods, evidence, analysis, ethics, language, formatting, and paragraph-level argumentation.
6. Do not treat "the paper has a basic thesis structure" as a merit.
7. A score around 70 means only passable. Do not inflate scores.
8. Keep the overall score consistent with the dimension scores and major issues.
9. If methods, evidence, or analysis are weak, the score should normally stay below 80.
10. Watch out for AI-sounding paragraphs that pile up academic terms but do not make a concrete argument.
11. If the paper uses generic AI-like writing or weak originality signals, it must not receive a high score.
12. If the research question is unclear, or the background/literature review does not clearly serve a central research goal, you must flag that as a major issue.

Paper text:
{text}
""".strip()
        return self.request_json_completion(cfg, SYSTEM_PROMPT, user_prompt, timeout=240)

    def split_sentences(self, line: str):
        parts = re.split(r"(?<=[。！？；])\s*", line)
        return [compact_whitespace(part) for part in parts if compact_whitespace(part)]

    def is_heading_line(self, line: str) -> bool:
        patterns = (
            r"^第[一二三四五六七八九十百]+章",
            r"^[0-9]{1,2}(?:\.[0-9]{1,2}){0,3}\s*\S+",
            r"^[一二三四五六七八九十]+、\S+",
            r"^(摘要|Abstract|关键词|Key words|参考文献|致谢|附录)$",
        )
        return len(line) <= 80 and any(re.match(pattern, line) for pattern in patterns)

    def is_noise_line(self, line: str) -> bool:
        if len(line) < 4:
            return True
        if re.fullmatch(r"[-=_.0-9\s]+", line):
            return True
        if re.match(r"^(表|图)\s*[0-9一二三四五六七八九十]", line):
            return True
        return False

    def is_reference_like_sentence(self, sentence: str) -> bool:
        lowered = sentence.lower()
        if lowered.startswith("[") and "]" in lowered[:6]:
            return True
        if "doi" in lowered:
            return True
        if re.search(r"\b\d{4}\b", sentence) and sentence.count(",") >= 2 and len(sentence) < 140:
            return True
        return False

    def extract_review_units(self, text: str):
        units = []
        current_section = "全文"
        sentence_id = 1
        for raw_line in text.splitlines():
            line = compact_whitespace(raw_line)
            if not line:
                continue
            if self.is_heading_line(line):
                current_section = shorten_text(line, 80)
                continue
            if self.is_noise_line(line):
                continue
            for sentence in self.split_sentences(line):
                if len(sentence) < 12 or self.is_reference_like_sentence(sentence):
                    continue
                units.append(
                    {
                        "id": sentence_id,
                        "section": current_section,
                        "label": f"{current_section} / 句子{sentence_id}",
                        "text": sentence,
                    }
                )
                sentence_id += 1
        return units

    def chunk_review_units(self, units):
        chunks = []
        current = []
        current_chars = 0
        for item in units:
            row = f"[{item['id']}] {item['label']}: {item['text']}"
            if current and (len(current) >= SENTENCE_SCAN_BATCH_SIZE or current_chars + len(row) > SENTENCE_SCAN_MAX_CHARS):
                chunks.append(current)
                current = []
                current_chars = 0
            current.append(item)
            current_chars += len(row)
        if current:
            chunks.append(current)
        return chunks

    def extract_paragraph_units(self, text: str):
        units = []
        current_section = "全文"
        paragraph_lines = []
        paragraph_id = 1

        def flush():
            nonlocal paragraph_lines, paragraph_id
            paragraph = compact_whitespace(" ".join(paragraph_lines))
            paragraph_lines = []
            if len(paragraph) < 30:
                return
            if self.is_reference_like_sentence(paragraph):
                return
            units.append(
                {
                    "id": paragraph_id,
                    "section": current_section,
                    "label": f"{current_section} / 段落{paragraph_id}",
                    "text": paragraph,
                }
            )
            paragraph_id += 1

        for raw_line in text.splitlines():
            line = compact_whitespace(raw_line)
            if not line:
                flush()
                continue
            if self.is_heading_line(line):
                flush()
                current_section = shorten_text(line, 80)
                continue
            if self.is_noise_line(line):
                flush()
                continue
            paragraph_lines.append(line)
        flush()
        return units

    def chunk_paragraph_units(self, units):
        chunks = []
        current = []
        current_chars = 0
        for item in units:
            row = f"[{item['id']}] {item['label']}: {item['text']}"
            if current and (len(current) >= PARAGRAPH_SCAN_BATCH_SIZE or current_chars + len(row) > PARAGRAPH_SCAN_MAX_CHARS):
                chunks.append(current)
                current = []
                current_chars = 0
            current.append(item)
            current_chars += len(row)
        if current:
            chunks.append(current)
        return chunks

    def generate_paragraph_argumentation_reviews(self, title: str, text: str, cfg: Config):
        units = self.extract_paragraph_units(text)
        if not units:
            return []
        reviews = []
        chunks = self.chunk_paragraph_units(units)
        for idx, chunk in enumerate(chunks, start=1):
            self.log(f"Paragraph scan {idx}/{len(chunks)}...")
            chunk_text = "\n".join(f"[{item['id']}] {item['label']}: {item['text']}" for item in chunk)
            label_lookup = {compact_whitespace(item["text"]): item["label"] for item in chunk}
            user_prompt = f"""
Paper title: {title}

Task:
1. Review the numbered thesis paragraphs below.
2. Judge whether each paragraph has real argumentation ability.
3. Return only paragraphs that are weak, generic, jargon-heavy, or unsupported.
4. Copy the provided paragraph label into the "section" field.
5. Keep the original paragraph excerpt in "quote" as exactly as possible.
6. Explain whether the paragraph lacks a claim, lacks evidence, jumps in logic, or mainly stacks empty academic rhetoric.

Paragraph list:
{chunk_text}
""".strip()
            data = self.request_json_completion(cfg, PARAGRAPH_SCAN_PROMPT, user_prompt, timeout=180)
            for item in data.get("paragraph_reviews", []):
                if not isinstance(item, dict):
                    continue
                quote = compact_whitespace(item.get("quote", ""))
                section = compact_whitespace(item.get("section", "")) or label_lookup.get(quote, "未定位")
                reviews.append(
                    {
                        "section": section,
                        "quote": quote,
                        "severity": item.get("severity", "medium"),
                        "weakness_type": item.get("weakness_type", "empty_rhetoric"),
                        "comment": item.get("comment", ""),
                        "suggestion": item.get("suggestion", ""),
                        "suggested_revision": item.get("suggested_revision", ""),
                    }
                )
        return dedupe_paragraph_reviews(reviews)

    def generate_sentence_level_annotations(self, title: str, text: str, cfg: Config):
        units = self.extract_review_units(text)
        if not units:
            return []
        annotations = []
        chunks = self.chunk_review_units(units)
        for idx, chunk in enumerate(chunks, start=1):
            self.log(f"Sentence scan {idx}/{len(chunks)}...")
            chunk_text = "\n".join(f"[{item['id']}] {item['label']}: {item['text']}" for item in chunk)
            label_lookup = {compact_whitespace(item["text"]): item["label"] for item in chunk}
            user_prompt = f"""
Paper title: {title}

Task:
1. Review the numbered thesis sentences below.
2. Identify every sentence that clearly needs correction or caution.
3. Skip sentences that are acceptable.
4. Copy the provided section label into the "section" field.
5. Keep the original sentence in "quote" exactly as shown when possible.
6. Give concise, local, actionable revision advice.

Sentence list:
{chunk_text}
""".strip()
            data = self.request_json_completion(cfg, SENTENCE_SCAN_PROMPT, user_prompt, timeout=180)
            for item in data.get("annotations", []):
                if not isinstance(item, dict):
                    continue
                quote = compact_whitespace(item.get("quote", ""))
                section = compact_whitespace(item.get("section", "")) or label_lookup.get(quote, "未定位")
                annotations.append(
                    {
                        "section": section,
                        "quote": quote,
                        "severity": item.get("severity", "medium"),
                        "issue_type": item.get("issue_type", "analysis"),
                        "comment": item.get("comment", ""),
                        "suggestion": item.get("suggestion", ""),
                        "suggested_revision": item.get("suggested_revision", ""),
                    }
                )
        return self.refine_generic_sentence_annotations(dedupe_annotations(annotations))

    def build_specific_language_feedback(self, sentence: str):
        jargon_hits = sum(1 for term in AI_JARGON_TERMS if term in sentence)
        value_hits = sum(1 for phrase in EMPTY_VALUE_PHRASES if phrase in sentence)
        if jargon_hits >= 3:
            return (
                "该句抽象名词和专业术语连续堆叠，但缺少关键动作、对象或关系说明，读者不容易把握作者到底想证明什么。",
                "减少术语串联，补出主语、对象和具体关系，把抽象概念改写成可理解的判断。",
            )
        if sentence.count("、") >= 4:
            return (
                "该句并列成分过多，像在连续列点，但各点之间的主次关系和逻辑关系没有交代清楚。",
                "保留最关键的两到三项，并说明这些内容之间是什么关系，不要只做罗列。",
            )
        if any(marker in sentence for marker in ("因此", "由此可见", "表明", "说明")):
            return (
                "该句把结论、依据和解释压在同一句里，容易造成论证跳步，读者看不清结论是如何得出的。",
                "把“依据是什么”和“据此得出什么结论”拆开写，先交代材料，再给出判断。",
            )
        if value_hits >= 1:
            return (
                "该句带有较强的价值判断色彩，但缺少支撑这些判断的具体内容，容易显得空泛。",
                "把“意义、价值、启示”落到具体对象、条件和机制上，再说明为什么成立。",
            )
        if len(sentence) >= 110 or sentence.count("，") >= 6:
            return (
                "该句包含的信息层次过多，主干判断被多重补充成分打断，影响读者快速抓住重点。",
                "把主结论单独写成一句，再把限定条件或补充解释另起一句展开。",
            )
        return (
            "该句虽能表达基本意思，但表述仍偏概括，学术写法还可以更具体。",
            "建议明确对象、范围或依据，避免停留在笼统判断。",
        )

    def refine_generic_sentence_annotations(self, annotations):
        refined = []
        generic_markers = (
            "该句过长，信息点叠加较多",
            "信息较密集，句子偏长",
        )
        for item in annotations:
            updated = dict(item)
            comment = compact_whitespace(updated.get("comment", ""))
            if updated.get("issue_type") == "language" and any(marker in comment for marker in generic_markers):
                better_comment, better_suggestion = self.build_specific_language_feedback(updated.get("quote", ""))
                updated["comment"] = better_comment
                if not compact_whitespace(updated.get("suggestion", "")) or "拆成" in compact_whitespace(updated.get("suggestion", "")):
                    updated["suggestion"] = better_suggestion
            refined.append(updated)
        return dedupe_annotations(refined)

    def demo_review(self, title: str, text: str):
        score = 74
        strengths = []
        issues = []
        paragraph_reviews = self.heuristic_paragraph_reviews(text)
        annotations = self.heuristic_sentence_annotations(text)

        if "社会工作" in text:
            strengths.append("选题具有社会工作专业视角与现实关怀。")
            score += 2
        if any(token in text for token in ("问卷", "访谈", "调查对象", "样本")):
            strengths.append("论文尝试使用经验材料支撑分析，而非完全停留在概念阐述。")
            score += 1
        if any(token in text for token in ("建议", "对策", "干预", "服务")) and "社会工作" in text:
            strengths.append("作者能够尝试把研究发现延伸到社会工作实践建议。")
            score += 1
        if "参考文献" in text:
            strengths.append("作者具备一定的文献搜集与整理基础。")
            score += 1

        if "链式中介模型成立" in text and "相关检验" in text:
            score -= 12
            issues.append("摘要或结论中的模型化表述超出了现有结果支持范围。")
            annotations.append(
                {
                    "section": "摘要 / 结论",
                    "quote": "链式中介模型成立",
                    "severity": "high",
                    "issue_type": "evidence",
                    "comment": "文中使用了较强的模型成立表述，但方法与结果未充分展示中介检验。",
                    "suggestion": "若当前仅有相关分析，应下调摘要和结论中的表述强度；若保留，则需补充中介分析结果。",
                    "suggested_revision": "研究结果初步提示社会支持不足可能通过负面情绪影响问题性网络使用，相关机制仍有待进一步统计检验。",
                }
            )

        tool_slice = text[text.find("3.4 研究工具") : text.find("3.5 数据处理")] if "3.4 研究工具" in text and "3.5 数据处理" in text else ""
        if "抑郁情绪" in text and "抑郁" not in tool_slice:
            score -= 8
            issues.append("研究变量与研究工具之间可能存在对应不完整的问题。")
            annotations.append(
                {
                    "section": "研究设计",
                    "quote": "3.4 研究工具",
                    "severity": "high",
                    "issue_type": "method",
                    "comment": "论文涉及多个核心变量，但工具说明未必全部完整呈现。",
                    "suggestion": "补充同伴支持和抑郁情绪对应的量表、题项数、评分方式与信效度。",
                    "suggested_revision": "",
                }
            )

        if "仅统计了母亲相关方面" in text:
            score -= 6
            issues.append("‘家庭支持’与‘仅统计母亲维度’之间存在概念口径不一致。")

        if "预计发放问卷500份" in text and "404" in text:
            score -= 4
            issues.append("方法部分仍停留在计划样本，未完全改成实际样本。")

        m = re.search(r"关键词[:：](.+)", text)
        if m:
            parts = [x.strip() for x in re.split(r"[；;，,]", m.group(1)) if x.strip()]
            if len(parts) < 4:
                score -= 3
                issues.append("关键词数量偏少，未完全覆盖研究核心变量。")

        if "[16] Kumar, Davis" in text:
            score -= 2
            issues.append("参考文献中存在作者格式异常，建议逐条核对。")

        if len(paragraph_reviews) >= 2:
            issues.append("部分段落存在术语堆砌、价值判断泛化或论点与依据脱节的问题，段落论证力不足。")
        if any(item["severity"] == "high" for item in paragraph_reviews):
            score -= 4
        elif paragraph_reviews:
            score -= 2

        if not strengths:
            strengths = ["论文围绕较为明确的研究主题展开，至少具备继续修改完善的基础。"]
        if not issues:
            issues = ["论文整体基础一般，仍应重点强化方法、结果与结论之间的对应关系。"]

        overall = (
            "该论文选题具有一定现实意义，也能从社会工作视角切入问题，说明作者具备基本的问题意识。"
            "但从目前稿件质量看，论文仍更接近“达到基础写作要求、尚需较大幅度修改”的水平，评分宜控制在合格到中等之间，"
            "不宜因篇章结构齐全而给予过高评价。论文的主要短板在于研究设计、统计处理与结论表达之间尚未充分对齐，"
            "部分判断存在表述强于证据、概念口径前后不够一致、工具说明不完整或结果支撑不足等问题。"
            "另外，若干句子的学术表达仍偏笼统，存在句子过长、措辞偏满、修改方向不够具体等现象，影响论证的严谨度。"
            "建议作者优先围绕变量界定、研究工具、统计分析、结论强度和参考文献规范进行实质性修订，"
            "并逐句打磨摘要、结果和结论中的关键表述。若能把证据链、方法说明和论证语言进一步校准，论文质量仍有明显提升空间。"
        )

        dims = [
            {"name": "选题与价值", "score": clamp_score(min(score + 3, 82)), "comment": "选题具有一定现实意义，但选题价值不能替代论证质量。"},
            {"name": "文献与理论", "score": clamp_score(score - 1), "comment": "已有一定文献与理论基础，但整合和提炼仍可加强。"},
            {"name": "方法与数据", "score": clamp_score(score - 6), "comment": "变量测量、样本描述与分析路径仍是主要薄弱环节。"},
            {"name": "分析与论证", "score": clamp_score(score - 5), "comment": "结果、证据与结论的对应关系仍需明显收紧。"},
            {"name": "社会工作实践性", "score": clamp_score(min(score + 1, 80)), "comment": "能提出实践建议，但建议需更直接回应研究发现。"},
            {"name": "语言与格式", "score": clamp_score(score - 2), "comment": "语言基本可读，但学术表达和格式规范性仍需细修。"},
        ]
        return {
            "title": title,
            "score": clamp_score(score),
            "overall_comment_zh": overall,
            "summary_strengths": strengths[:5],
            "major_issues": issues[:6],
            "dimension_scores": dims,
            "annotations": dedupe_annotations(annotations, limit=40),
            "paragraph_reviews": dedupe_paragraph_reviews(paragraph_reviews, limit=20),
        }

    def heuristic_sentence_annotations(self, text: str):
        annotations = []
        units = self.extract_review_units(text)
        for item in units:
            sentence = item["text"]
            if "链式中介模型成立" in sentence or re.search(r"(充分证明|完全证明|必然导致|显著说明)", sentence):
                annotations.append(
                    {
                        "section": item["label"],
                        "quote": sentence,
                        "severity": "high",
                        "issue_type": "evidence",
                        "comment": "该句使用了过强的结论化措辞，但从句面上看未同步呈现足够证据。",
                        "suggestion": "把“成立、证明、必然”等结论性表述改为“提示、可能、初步发现”等更审慎表述。",
                        "suggested_revision": "",
                    }
                )
            elif "显著" in sentence and not re.search(r"(p\s*[<=>]|P\s*[<=>]|r\s*=|t\s*=|F\s*=|β\s*=|\d+\.\d+)", sentence):
                annotations.append(
                    {
                        "section": item["label"],
                        "quote": sentence,
                        "severity": "medium",
                        "issue_type": "evidence",
                        "comment": "该句提到显著性或明显差异，但未同步给出统计依据，容易显得证据不足。",
                        "suggestion": "补充相应统计指标，或改成不超过现有结果的描述。",
                        "suggested_revision": "",
                    }
                )
            elif len(sentence) >= 110 or sentence.count("，") >= 6:
                comment, suggestion = self.build_specific_language_feedback(sentence)
                annotations.append(
                    {
                        "section": item["label"],
                        "quote": sentence,
                        "severity": "medium",
                        "issue_type": "language",
                        "comment": comment,
                        "suggestion": suggestion,
                        "suggested_revision": "",
                    }
                )
            elif re.search(r"(一定程度上|较为|比较|很多|一些|相关方面|等等|等方面)", sentence):
                annotations.append(
                    {
                        "section": item["label"],
                        "quote": sentence,
                        "severity": "low",
                        "issue_type": "language",
                        "comment": "该句措辞偏笼统，学术表达不够精准，容易削弱论证力度。",
                        "suggestion": "尽量替换成更可检验、更可界定的具体表达，并说明对象或范围。",
                        "suggested_revision": "",
                    }
                )
            elif sentence.startswith("因此") and re.search(r"(说明|证明|表明)", sentence):
                annotations.append(
                    {
                        "section": item["label"],
                        "quote": sentence,
                        "severity": "medium",
                        "issue_type": "analysis",
                        "comment": "该句从前文结果直接推到较强结论，存在推断跨度偏大的风险。",
                        "suggestion": "补充中间论证环节，或改成较为审慎的解释性表述。",
                        "suggested_revision": "",
                    }
                )
        return self.refine_generic_sentence_annotations(dedupe_annotations(annotations, limit=40))

    def heuristic_paragraph_reviews(self, text: str):
        reviews = []
        strong_evidence_markers = ("数据", "结果显示", "结果表明", "结果发现", "访谈", "问卷", "调查", "案例", "统计", "相关系数", "回归", "显著", "样本")
        weak_connector_markers = ("因此", "由此可见", "说明", "表明", "可见")
        for item in self.extract_paragraph_units(text):
            paragraph = item["text"]
            jargon_hits = sum(1 for term in AI_JARGON_TERMS if term in paragraph)
            evidence_hits = sum(1 for marker in strong_evidence_markers if marker in paragraph)
            connector_hits = sum(1 for marker in weak_connector_markers if marker in paragraph)
            value_hits = sum(1 for phrase in EMPTY_VALUE_PHRASES if phrase in paragraph)
            has_number = bool(re.search(r"\d", paragraph))
            has_citation = bool(re.search(r"[\[\(（]\d+[\]\)）]", paragraph))
            has_claim_jump = any(marker in paragraph for marker in ("因此", "由此可见", "这说明", "表明了"))
            looks_intro_paragraph = item["id"] <= 3 or any(hint in item["section"] for hint in INTRO_SECTION_HINTS)
            has_research_question_marker = any(marker in paragraph for marker in RESEARCH_QUESTION_MARKERS)

            has_real_support = evidence_hits > 0 or has_number or has_citation

            if (jargon_hits >= 5 or (jargon_hits >= 4 and value_hits >= 1)) and not has_real_support:
                reviews.append(
                    {
                        "section": item["label"],
                        "quote": shorten_text(paragraph, 200),
                        "severity": "high",
                        "weakness_type": "ai_generic_language",
                        "comment": "这一段很像 AI 生成的通用学术话术：术语和价值判断很多，但缺少明确对象、材料和原创分析。",
                        "suggestion": "删除模板化表述，改成作者自己的具体判断，并补充与研究主题直接相关的事实、文献或数据支撑。",
                        "suggested_revision": "",
                    }
                )
            elif jargon_hits >= 4 and not has_real_support:
                reviews.append(
                    {
                        "section": item["label"],
                        "quote": shorten_text(paragraph, 200),
                        "severity": "high",
                        "weakness_type": "jargon_stack",
                        "comment": "这一段术语密度较高，但缺少可核实的对象、证据或展开说明，容易形成“像在论证、其实没有论证”的空转表达。",
                        "suggestion": "先明确本段核心判断，再补充事实、数据、文本证据或具体解释，减少空泛术语串联。",
                        "suggested_revision": "",
                    }
                )
            elif looks_intro_paragraph and not has_research_question_marker and value_hits >= 1 and not has_real_support:
                reviews.append(
                    {
                        "section": item["label"],
                        "quote": shorten_text(paragraph, 200),
                        "severity": "high" if item["id"] <= 2 else "medium",
                        "weakness_type": "research_question_blur",
                        "comment": "这一段没有把研究问题、研究对象或核心变量说清楚，导致背景论述和后续综述难以服务一个明确目标。",
                        "suggestion": "直接写明本文到底研究什么问题、聚焦哪些对象和变量，再围绕这个目标重写背景与综述。",
                        "suggested_revision": "",
                    }
                )
            elif value_hits >= 2 and not has_real_support:
                reviews.append(
                    {
                        "section": item["label"],
                        "quote": shorten_text(paragraph, 200),
                        "severity": "medium",
                        "weakness_type": "empty_rhetoric",
                        "comment": "这一段主要停留在“有意义、有价值、有助于”等泛化判断上，没有形成充分论证。",
                        "suggestion": "把价值判断拆开，说明具体对谁、在什么条件下、通过什么机制产生作用。",
                        "suggested_revision": "",
                    }
                )
            elif has_claim_jump and not has_real_support:
                reviews.append(
                    {
                        "section": item["label"],
                        "quote": shorten_text(paragraph, 200),
                        "severity": "medium",
                        "weakness_type": "logic_jump",
                        "comment": "这一段直接推出结论，但段内没有提供足够依据来支撑该判断，论证链条偏短。",
                        "suggestion": "补充前提、材料或解释步骤，再导出结论，避免“先下判断、后找理由”。",
                        "suggested_revision": "",
                    }
                )
            elif jargon_hits >= 2 and paragraph.count("、") >= 4 and evidence_hits == 0 and connector_hits <= 1 and not has_real_support:
                reviews.append(
                    {
                        "section": item["label"],
                        "quote": shorten_text(paragraph, 200),
                        "severity": "low",
                        "weakness_type": "no_evidence",
                        "comment": "这一段列举了多个抽象要点，但仍缺少支撑这些要点的解释或证据。",
                        "suggestion": "保留最关键的两到三点，并逐点说明依据，而不是连续罗列概念。",
                        "suggested_revision": "",
                    }
                )
        return dedupe_paragraph_reviews(reviews, limit=20)

    def read_with_word(self, path: Path):
        if not can_use_word_com():
            raise ReviewError("Microsoft Word automation is only available on Windows with pywin32 installed.")
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = create_word_application()
            doc = call_with_retries(
                word.Documents.Open,
                str(path.resolve()),
                False,
                True,
                retries=10,
                delay=0.6,
            )
            return call_with_retries(lambda: doc.Content.Text, retries=10, delay=0.5)
        except Exception as exc:
            raise ReviewError(f"Failed to read document with Word: {exc}") from exc
        finally:
            if doc is not None:
                try:
                    call_with_retries(doc.Close, False, retries=6, delay=0.3)
                except Exception:
                    pass
            if word is not None:
                try:
                    call_with_retries(word.Quit, retries=6, delay=0.3)
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    def write_outputs(self, source: Path, out_dir: Path, review: dict, make_annotated: bool):
        out_dir.mkdir(parents=True, exist_ok=True)
        slug = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = safe_name(source.stem)
        json_path = out_dir / f"{stem}_review_{slug}.json"
        json_path.write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")

        self.log("Generating detailed annotation document...")
        report_path = out_dir / f"{stem}_详细批注_{slug}.docx"
        self.make_report_docx(report_path, source.name, review)

        self.log("Generating summary and mark document...")
        summary_path = out_dir / f"{stem}_评分与总结_{slug}.docx"
        self.make_summary_docx(summary_path, source.name, review)

        annotated_path = ""
        if make_annotated and source.suffix.lower() in {".doc", ".docx", ".rtf"}:
            if not can_use_word_com():
                self.log("Annotated copy skipped: inline Word comments are only available on Windows with Microsoft Word.")
            else:
                try:
                    self.log("Generating annotated copy...")
                    annotated = out_dir / f"{stem}_annotated_{slug}.docx"
                    self.make_annotated_doc(source, annotated, review)
                    annotated_path = str(annotated)
                except Exception as exc:
                    self.log(f"Annotated copy skipped: {exc}")

        return {
            "json": str(json_path),
            "detailed_docx": str(report_path),
            "summary_docx": str(summary_path),
            "annotated_docx": annotated_path,
        }

    def make_report_docx(self, path: Path, source_name: str, review: dict):
        if DocxDocument is not None:
            try:
                doc = DocxDocument()
                doc.add_heading("论文详细批注报告", level=0)
                for line in (
                    f"原始文件：{source_name}",
                    f"论文标题：{review['title']}",
                    f"评审模式：{review.get('review_mode', 'unknown')}",
                    f"评分：{review['score']} / 100",
                    f"段落论证问题：{len(review.get('paragraph_reviews', []))}",
                    f"逐句批注条数：{len(review['annotations'])}",
                    f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                ):
                    doc.add_paragraph(line)

                doc.add_heading("使用说明", level=1)
                doc.add_paragraph("本报告优先指出与研究问题、概念界定、文献梳理、研究方法、证据使用、结果分析和社会工作专业逻辑相关的实质性问题；仅在必要时补充语言或格式问题。")

                doc.add_heading("段落论证力诊断", level=1)
                if review.get("paragraph_reviews"):
                    for idx, item in enumerate(review.get("paragraph_reviews", []), start=1):
                        doc.add_paragraph(f"段落诊断 {idx}")
                        doc.add_paragraph(f"位置：{item['section']}")
                        doc.add_paragraph(f"段落摘录：{item['quote']}")
                        doc.add_paragraph(f"问题等级：{item['severity']}    类型：{item['weakness_type']}")
                        doc.add_paragraph(f"诊断：{item['comment']}")
                        doc.add_paragraph(f"修改建议：{item['suggestion']}")
                        if item["suggested_revision"]:
                            doc.add_paragraph(f"参考改写：{item['suggested_revision']}")
                else:
                    doc.add_paragraph("未检出显著的段落级空泛论证问题。")

                doc.add_heading("主要问题", level=1)
                for item in review["major_issues"] or ["未提供。"]:
                    doc.add_paragraph(f"- {item}")

                doc.add_heading("逐句详细批注", level=1)
                if not review["annotations"]:
                    doc.add_paragraph("未生成逐句批注。")
                for idx, item in enumerate(review["annotations"], start=1):
                    doc.add_paragraph(f"批注 {idx}")
                    doc.add_paragraph(f"位置：{item['section']}")
                    doc.add_paragraph(f"原句：{item['quote']}")
                    doc.add_paragraph(f"问题等级：{item['severity']}    问题类型：{item['issue_type']}")
                    doc.add_paragraph(f"问题说明：{item['comment']}")
                    doc.add_paragraph(f"修改建议：{item['suggestion']}")
                    if item["suggested_revision"]:
                        doc.add_paragraph(f"参考改写：{item['suggested_revision']}")

                path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(path.resolve()))
                return
            except Exception as exc:
                if not can_use_word_com():
                    raise ReviewError(f"Failed to create DOCX report: {exc}") from exc
        if not can_use_word_com():
            raise ReviewError("DOCX output requires python-docx, or Microsoft Word automation on Windows.")
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = create_word_application()
            word.Documents.Add()
            doc = word.ActiveDocument
            sel = word.Selection
            sel.Font.Name = "Microsoft YaHei"
            sel.Font.Size = 18
            sel.Font.Bold = True
            sel.TypeText("论文段落与逐句问题批注报告")
            sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 11
            sel.Font.Bold = False
            for line in (
                f"原始文件：{source_name}",
                f"论文标题：{review['title']}",
                f"评审模式：{review.get('review_mode', 'unknown')}",
                f"评分：{review['score']} / 100",
                f"段落论证问题：{len(review.get('paragraph_reviews', []))}",
                f"逐句批注条数：{len(review['annotations'])}",
                f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            ):
                sel.TypeText(line)
                sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("使用说明")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            sel.TypeText("本报告重点列出需要优先修改的句子或短语，并逐条说明问题、修改方向和可选改写。未被列出的句子不等于完全没有问题，但建议优先处理下列内容。")
            sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("段落论证力诊断")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            if not review.get("paragraph_reviews"):
                sel.TypeText("未检出明显的段落级空泛论证问题。")
                sel.TypeParagraph()
            for idx, item in enumerate(review.get("paragraph_reviews", []), start=1):
                sel.Font.Bold = True
                sel.TypeText(f"段落诊断 {idx}")
                sel.TypeParagraph()
                sel.Font.Bold = False
                sel.TypeText(f"位置：{item['section']}")
                sel.TypeParagraph()
                sel.TypeText(f"段落摘录：{item['quote']}")
                sel.TypeParagraph()
                sel.TypeText(f"问题等级：{item['severity']}    薄弱类型：{item['weakness_type']}")
                sel.TypeParagraph()
                sel.TypeText(f"诊断：{item['comment']}")
                sel.TypeParagraph()
                sel.TypeText(f"修改建议：{item['suggestion']}")
                sel.TypeParagraph()
                if item["suggested_revision"]:
                    sel.TypeText(f"参考改写：{item['suggested_revision']}")
                    sel.TypeParagraph()
                sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("优先问题")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            for item in review["major_issues"] or ["未提供。"]:
                sel.TypeText(f"• {item}")
                sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("逐句问题批注")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            if not review["annotations"]:
                sel.TypeText("未生成逐句批注。")
                sel.TypeParagraph()
            for idx, item in enumerate(review["annotations"], start=1):
                sel.Font.Bold = True
                sel.TypeText(f"批注 {idx}")
                sel.TypeParagraph()
                sel.Font.Bold = False
                sel.TypeText(f"位置：{item['section']}")
                sel.TypeParagraph()
                sel.TypeText(f"原句：{item['quote']}")
                sel.TypeParagraph()
                sel.TypeText(f"问题等级：{item['severity']}    问题类型：{item['issue_type']}")
                sel.TypeParagraph()
                sel.TypeText(f"问题说明：{item['comment']}")
                sel.TypeParagraph()
                sel.TypeText(f"修改建议：{item['suggestion']}")
                sel.TypeParagraph()
                if item["suggested_revision"]:
                    sel.TypeText(f"参考改写：{item['suggested_revision']}")
                    sel.TypeParagraph()
                sel.TypeParagraph()

            time.sleep(0.5)
            call_with_retries(doc.SaveAs2, str(path.resolve()), FileFormat=16)
            time.sleep(0.3)
        except Exception as exc:
            raise ReviewError(f"Failed to create Word report: {exc}") from exc
        finally:
            if doc is not None:
                try:
                    call_with_retries(doc.Close, False, retries=6, delay=0.3)
                except Exception:
                    pass
            if word is not None:
                try:
                    call_with_retries(word.Quit, retries=6, delay=0.3)
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    def make_summary_docx(self, path: Path, source_name: str, review: dict):
        if DocxDocument is not None:
            try:
                doc = DocxDocument()
                doc.add_heading("评分与总体总结", level=0)
                for line in (
                    f"原始文件：{source_name}",
                    f"论文标题：{review['title']}",
                    f"评审模式：{review.get('review_mode', 'unknown')}",
                    f"评分：{review['score']} / 100",
                    f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                ):
                    doc.add_paragraph(line)
                if review.get("score_calibration_note"):
                    doc.add_paragraph(f"评分校准说明：{review['score_calibration_note']}")
                doc.add_paragraph(f"段落论证问题：{len(review.get('paragraph_reviews', []))} 段")

                doc.add_heading("300字左右总体评语", level=1)
                doc.add_paragraph(review["overall_comment_zh"])

                doc.add_heading("段落论证提醒", level=1)
                if review.get("paragraph_reviews"):
                    for item in review["paragraph_reviews"][:6]:
                        doc.add_paragraph(f"- {item['section']}：{item['comment']}")
                else:
                    doc.add_paragraph("未检出显著的段落级空泛论证问题。")

                doc.add_heading("核心优点", level=1)
                for item in review["summary_strengths"] or ["未提供。"]:
                    doc.add_paragraph(f"- {item}")

                doc.add_heading("主要问题", level=1)
                for item in review["major_issues"] or ["未提供。"]:
                    doc.add_paragraph(f"- {item}")

                doc.add_heading("评分维度", level=1)
                for item in review["dimension_scores"]:
                    doc.add_paragraph(f"{item['name']}：{item['score']} 分")
                    if item["comment"]:
                        doc.add_paragraph(f"说明：{item['comment']}")

                path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(path.resolve()))
                return
            except Exception as exc:
                if not can_use_word_com():
                    raise ReviewError(f"Failed to create summary document: {exc}") from exc
        if not can_use_word_com():
            raise ReviewError("DOCX output requires python-docx, or Microsoft Word automation on Windows.")
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = create_word_application()
            word.Documents.Add()
            doc = word.ActiveDocument
            sel = word.Selection
            sel.Font.Name = "Microsoft YaHei"
            sel.Font.Size = 18
            sel.Font.Bold = True
            sel.TypeText("评分与300字总结")
            sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 11
            sel.Font.Bold = False
            sel.TypeText(f"原始文件：{source_name}")
            sel.TypeParagraph()
            sel.TypeText(f"论文标题：{review['title']}")
            sel.TypeParagraph()
            sel.TypeText(f"评审模式：{review.get('review_mode', 'unknown')}")
            sel.TypeParagraph()
            sel.TypeText(f"评分：{review['score']} / 100")
            sel.TypeParagraph()
            sel.TypeText(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            sel.TypeParagraph()
            if review.get("score_calibration_note"):
                sel.TypeText(f"评分校准说明：{review['score_calibration_note']}")
                sel.TypeParagraph()
            sel.TypeText(f"段落论证问题：{len(review.get('paragraph_reviews', []))} 段")
            sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("总体批注总结")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            sel.TypeText(review["overall_comment_zh"])
            sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("段落论证力提醒")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            if review.get("paragraph_reviews"):
                for item in review["paragraph_reviews"][:6]:
                    sel.TypeText(f"• {item['section']}：{item['comment']}")
                    sel.TypeParagraph()
            else:
                sel.TypeText("未检出明显的段落级空泛论证问题。")
                sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("核心优点")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            for item in review["summary_strengths"] or ["未提供。"]:
                sel.TypeText(f"• {item}")
                sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("主要问题")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            for item in review["major_issues"] or ["未提供。"]:
                sel.TypeText(f"• {item}")
                sel.TypeParagraph()
            sel.TypeParagraph()

            sel.Font.Size = 14
            sel.Font.Bold = True
            sel.TypeText("评分维度")
            sel.TypeParagraph()
            sel.Font.Size = 11
            sel.Font.Bold = False
            for item in review["dimension_scores"]:
                sel.TypeText(f"{item['name']}：{item['score']} 分")
                sel.TypeParagraph()
                if item["comment"]:
                    sel.TypeText(f"说明：{item['comment']}")
                    sel.TypeParagraph()
                sel.TypeParagraph()

            time.sleep(0.5)
            call_with_retries(doc.SaveAs2, str(path.resolve()), FileFormat=16)
            time.sleep(0.3)
        except Exception as exc:
            raise ReviewError(f"Failed to create summary document: {exc}") from exc
        finally:
            if doc is not None:
                try:
                    call_with_retries(doc.Close, False, retries=6, delay=0.3)
                except Exception:
                    pass
            if word is not None:
                try:
                    call_with_retries(word.Quit, retries=6, delay=0.3)
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    def make_annotated_doc(self, source: Path, target: Path, review: dict):
        if not can_use_word_com():
            raise ReviewError("Inline Word comments are only available on Windows with Microsoft Word.")
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = create_word_application()
            doc = call_with_retries(word.Documents.Open, str(source.resolve()), retries=10, delay=0.6)
            for item in review["annotations"]:
                quote = (item.get("quote") or "").strip()
                if not quote:
                    continue
                text = f"[{item['severity']}/{item['issue_type']}] {item['comment']}\n建议：{item['suggestion']}"
                rng = doc.Content
                find = rng.Find
                find.ClearFormatting()
                find.Text = quote[:100]
                if find.Execute():
                    doc.Comments.Add(rng, text[:1500])
            time.sleep(0.5)
            call_with_retries(doc.SaveAs2, str(target.resolve()), FileFormat=16)
        except Exception as exc:
            raise ReviewError(f"Failed to create annotated document: {exc}") from exc
        finally:
            if doc is not None:
                try:
                    call_with_retries(doc.Close, False, retries=6, delay=0.3)
                except Exception:
                    pass
            if word is not None:
                try:
                    call_with_retries(word.Quit, retries=6, delay=0.3)
                except Exception:
                    pass
            pythoncom.CoUninitialize()


class App:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title(APP_TITLE)
        self.root.geometry("980x760")
        self.root.minsize(920, 700)

        self.paper_var = tk.StringVar()
        self.output_var = tk.StringVar(value=str((Path.cwd() / "thesis_reviewer_output").resolve()))
        self.api_base_var = tk.StringVar(value="https://api.openai.com/v1")
        self.api_key_var = tk.StringVar()
        self.model_var = tk.StringVar(value="gpt-4.1")
        self.demo_var = tk.BooleanVar(value=False)
        self.annotated_var = tk.BooleanVar(value=IS_WINDOWS)
        self.score_var = tk.StringVar(value="--")
        self.status_var = tk.StringVar(value="Ready")
        self.queue = queue.Queue()
        self.worker = None

        self.build_ui()
        self.load_settings()
        self.configure_platform_capabilities()
        self.root.after(150, self.poll_queue)

    def build_ui(self):
        style = ttk.Style(self.root)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        wrap = ttk.Frame(self.root, padding=16)
        wrap.pack(fill="both", expand=True)
        wrap.columnconfigure(0, weight=1)
        wrap.rowconfigure(3, weight=1)

        ttk.Label(wrap, text=APP_TITLE, font=("Segoe UI", 18, "bold")).grid(row=0, column=0, sticky="w")
        ttk.Label(
            wrap,
            text="Upload a thesis, click 执行评审, and generate a mark, a 300-word overall comment, and a sentence-level Word review report.",
            wraplength=900,
        ).grid(row=1, column=0, sticky="we", pady=(6, 14))

        form = ttk.LabelFrame(wrap, text="Settings / 设置", padding=12)
        form.grid(row=2, column=0, sticky="ew")
        form.columnconfigure(1, weight=1)

        ttk.Label(form, text="Paper / 论文文件").grid(row=0, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.paper_var).grid(row=0, column=1, sticky="ew", padx=8)
        ttk.Button(form, text="Browse", command=self.pick_paper).grid(row=0, column=2)

        ttk.Label(form, text="Output / 输出目录").grid(row=1, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.output_var).grid(row=1, column=1, sticky="ew", padx=8)
        ttk.Button(form, text="Browse", command=self.pick_output).grid(row=1, column=2)

        ttk.Label(form, text="API Base URL").grid(row=2, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.api_base_var).grid(row=2, column=1, sticky="ew", padx=8)

        ttk.Label(form, text="Model").grid(row=3, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.model_var).grid(row=3, column=1, sticky="ew", padx=8)

        ttk.Label(form, text="API Key").grid(row=4, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.api_key_var, show="*").grid(row=4, column=1, sticky="ew", padx=8)

        ttk.Checkbutton(
            form,
            text="Use demo heuristic review if API is unavailable / 无 API 时使用演示评审",
            variable=self.demo_var,
        ).grid(row=5, column=0, columnspan=3, sticky="w", pady=(8, 4))

        self.annotated_check = ttk.Checkbutton(
            form,
            text="Create annotated Word copy for .doc/.docx / 生成 Word 批注版",
            variable=self.annotated_var,
        )
        self.annotated_check.grid(row=6, column=0, columnspan=3, sticky="w")

        actions = ttk.Frame(form)
        actions.grid(row=7, column=0, columnspan=3, sticky="ew", pady=(10, 0))
        actions.columnconfigure(1, weight=1)
        ttk.Button(actions, text="执行评审 / Run Review", command=self.start_review).grid(row=0, column=0, sticky="w")
        ttk.Label(actions, text="Mark").grid(row=0, column=1, sticky="e", padx=(0, 8))
        ttk.Label(actions, textvariable=self.score_var, font=("Segoe UI", 16, "bold")).grid(row=0, column=2, sticky="e")

        panes = ttk.PanedWindow(wrap, orient="vertical")
        panes.grid(row=3, column=0, sticky="nsew", pady=(14, 0))

        result = ttk.LabelFrame(panes, text="Result / 结果", padding=12)
        logs = ttk.LabelFrame(panes, text="Log / 日志", padding=12)
        panes.add(result, weight=3)
        panes.add(logs, weight=2)

        result.columnconfigure(1, weight=1)
        self.comment_box = tk.Text(result, wrap="word", height=18)
        self.comment_box.grid(row=0, column=0, columnspan=2, sticky="nsew")
        self.comment_box.configure(state="disabled")

        self.detailed_var = tk.StringVar()
        self.summary_path_var = tk.StringVar()
        ttk.Label(result, text="Detailed Sentence Notes Doc").grid(row=1, column=0, sticky="w", pady=(10, 0))
        ttk.Entry(result, textvariable=self.detailed_var).grid(row=1, column=1, sticky="ew", pady=(10, 0))
        ttk.Label(result, text="Summary + Mark Doc").grid(row=2, column=0, sticky="w", pady=(6, 0))
        ttk.Entry(result, textvariable=self.summary_path_var).grid(row=2, column=1, sticky="ew", pady=(6, 0))
        ttk.Button(result, text="Open Output Folder", command=self.open_output).grid(row=3, column=1, sticky="e", pady=(10, 0))

        self.log_box = tk.Text(logs, wrap="word", height=12)
        self.log_box.pack(fill="both", expand=True)
        self.log_box.configure(state="disabled")

        status_bar = ttk.Frame(wrap)
        status_bar.grid(row=4, column=0, sticky="ew", pady=(10, 0))
        status_bar.columnconfigure(0, weight=1)
        ttk.Label(status_bar, textvariable=self.status_var).grid(row=0, column=0, sticky="w")

    def configure_platform_capabilities(self):
        if IS_WINDOWS:
            return
        self.annotated_var.set(False)
        try:
            self.annotated_check.state(["disabled"])
        except Exception:
            pass
        self.append_log("Inline Word comments are disabled on this platform. The detailed review doc and summary doc still work.")

    def pick_paper(self):
        path = filedialog.askopenfilename(
            title="Select thesis paper",
            filetypes=[
                ("Supported files", "*.doc *.docx *.txt *.md *.rtf *.pdf"),
                ("All files", "*.*"),
            ],
        )
        if path:
            self.paper_var.set(path)

    def pick_output(self):
        path = filedialog.askdirectory(title="Select output folder")
        if path:
            self.output_var.set(path)

    def append_log(self, message: str):
        self.log_box.configure(state="normal")
        self.log_box.insert("end", f"[{datetime.now().strftime('%H:%M:%S')}] {message}\n")
        self.log_box.see("end")
        self.log_box.configure(state="disabled")

    def set_comment(self, text: str):
        self.comment_box.configure(state="normal")
        self.comment_box.delete("1.0", "end")
        self.comment_box.insert("1.0", text)
        self.comment_box.configure(state="disabled")

    def log(self, message: str):
        self.queue.put(("log", message))

    def start_review(self):
        if self.worker and self.worker.is_alive():
            messagebox.showinfo(APP_TITLE, "A review is already running.")
            return
        paper = Path(self.paper_var.get().strip())
        if not paper.exists():
            messagebox.showerror(APP_TITLE, "Please choose a valid paper file.")
            return
        cfg = Config(
            paper=paper,
            output_dir=Path(self.output_var.get().strip() or Path.cwd()),
            api_base=self.api_base_var.get().strip(),
            api_key=self.api_key_var.get().strip(),
            model=self.model_var.get().strip(),
            fallback_demo=self.demo_var.get(),
            annotated_copy=self.annotated_var.get(),
        )
        self.save_settings()
        self.score_var.set("--")
        self.detailed_var.set("")
        self.summary_path_var.set("")
        self.set_comment("")
        self.status_var.set("Running review...")
        self.append_log("Starting review...")

        def work():
            try:
                engine = ReviewEngine(self.log)
                result = engine.review(cfg)
                self.queue.put(("done", result))
            except Exception as exc:
                self.queue.put(("error", f"{exc}\n\n{traceback.format_exc()}"))

        self.worker = threading.Thread(target=work, daemon=True)
        self.worker.start()

    def poll_queue(self):
        try:
            while True:
                kind, payload = self.queue.get_nowait()
                if kind == "log":
                    self.append_log(str(payload))
                elif kind == "done":
                    self.score_var.set(str(payload.get("score", "--")))
                    self.set_comment(payload.get("overall_comment_zh", ""))
                    outputs = payload.get("outputs", {})
                    self.detailed_var.set(outputs.get("detailed_docx", ""))
                    self.summary_path_var.set(outputs.get("summary_docx", ""))
                    self.append_log("Review completed.")
                    self.status_var.set("Review completed.")
                elif kind == "error":
                    self.append_log("Review failed.")
                    self.status_var.set("Review failed.")
                    messagebox.showerror(APP_TITLE, str(payload))
        except queue.Empty:
            pass
        self.root.after(150, self.poll_queue)

    def open_output(self):
        target = self.detailed_var.get().strip() or self.summary_path_var.get().strip() or self.output_var.get().strip()
        if not target:
            return
        folder = Path(target)
        folder = folder if folder.is_dir() else folder.parent
        if folder.exists():
            try:
                open_path(folder)
            except Exception as exc:
                messagebox.showerror(APP_TITLE, f"Failed to open output folder: {exc}")

    def save_settings(self):
        data = {
            "output_dir": self.output_var.get().strip(),
            "api_base": self.api_base_var.get().strip(),
            "model": self.model_var.get().strip(),
            "fallback_demo": self.demo_var.get(),
            "annotated_copy": self.annotated_var.get(),
        }
        SETTINGS_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    def load_settings(self):
        if not SETTINGS_PATH.exists():
            return
        try:
            data = json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
        except Exception:
            return
        self.output_var.set(data.get("output_dir", self.output_var.get()))
        self.api_base_var.set(data.get("api_base", self.api_base_var.get()))
        self.model_var.set(data.get("model", self.model_var.get()))
        self.demo_var.set(bool(data.get("fallback_demo", self.demo_var.get())))
        self.annotated_var.set(bool(data.get("annotated_copy", self.annotated_var.get())))


def main():
    root = tk.Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
